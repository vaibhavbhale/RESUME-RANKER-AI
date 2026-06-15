from pypdf import PdfReader
import docx

def extract_text_from_pdf(path: str) -> str:
    reader = PdfReader(path)
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts).strip()

def extract_text_from_docx(path: str) -> str:
    d = docx.Document(path)
    return "\n".join([p.text for p in d.paragraphs]).strip()

def extract_text(resume) -> str:
    path = resume.file.path
    name = (resume.original_filename or "").lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf(path)
    if name.endswith(".docx"):
        return extract_text_from_docx(path)
    raise ValueError("Unsupported file type (PDF/DOCX only)")

def extract_text_from_upload(uploaded_file) -> str:
    name = (getattr(uploaded_file, "name", "") or "").lower()
    try:
        uploaded_file.seek(0)
    except Exception:
        pass

    if name.endswith(".pdf"):
        reader = PdfReader(uploaded_file)
        parts = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        return "\n".join(parts).strip()

    if name.endswith(".docx"):
        d = docx.Document(uploaded_file)
        return "\n".join([p.text for p in d.paragraphs]).strip()

    raise ValueError("Unsupported file type (PDF/DOCX only)")

from pypdf import PdfReader
import docx

def extract_text_from_upload(uploaded_file) -> str:
    """
    Extract text from an uploaded file object (InMemoryUploadedFile / TemporaryUploadedFile)
    WITHOUT saving to disk. Supports .pdf and .docx
    """
    name = (getattr(uploaded_file, "name", "") or "").lower()

    try:
        uploaded_file.seek(0)
    except Exception:
        pass

    if name.endswith(".pdf"):
        reader = PdfReader(uploaded_file)
        parts = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        return "\n".join(parts).strip()

    if name.endswith(".docx"):
        d = docx.Document(uploaded_file)
        return "\n".join(p.text for p in d.paragraphs).strip()

    raise ValueError("Unsupported file type. Only PDF/DOCX supported.")